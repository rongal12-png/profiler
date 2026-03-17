import math
import pandas as pd
from jinja2 import Environment, FileSystemLoader
import markdown2
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import io
import base64
from pathlib import Path
from datetime import datetime

from .chains import get_known_label
from . import project_health

# Setup Jinja2 environment (M10: autoescape=True to prevent XSS from blockchain data)
template_loader = FileSystemLoader(searchpath=Path(__file__).parent / "templates")
jinja_env = Environment(loader=template_loader, autoescape=True)


def generate_reference_id(job_id: int) -> str:
    """
    Generates a sophisticated reference ID from a job ID.
    Format: WI-{YYMMDD}-{JOB_ID:05d}
    Example: WI-260215-00042
    """
    date_part = datetime.now().strftime('%y%m%d')
    return f"WI-{date_part}-{job_id:05d}"


def results_to_dataframe(results: list) -> pd.DataFrame:
    """Converts a list of WalletAnalysis ORM objects to a clean pandas DataFrame."""
    data = [res.__dict__ for res in results]
    df = pd.DataFrame(data)

    # Clean up and format the dataframe
    drop_cols = [c for c in ['_sa_instance_state', 'id', 'job_id'] if c in df.columns]
    df = df.drop(columns=drop_cols)
    df['last_scored_at'] = pd.to_datetime(df['last_scored_at'])

    # Flatten nested JSON fields for CSV/analyst view
    df['stable_usd_total'] = df['stable_balances'].apply(lambda x: sum(float(item.get('usd', 0)) for item in x) if isinstance(x, list) and x else 0)
    df['labels_str'] = df['labels'].apply(lambda x: ', '.join(str(i) for i in x) if isinstance(x, list) and x else '')
    df['risk_flags_str'] = df['risk_flags'].apply(lambda x: ', '.join(str(i) for i in x) if isinstance(x, list) and x else '')
    df['tx_count'] = df['activity_indicators'].apply(lambda x: x.get('tx_count', 0) if isinstance(x, dict) else 0)

    # Ensure new scoring columns exist (backward compat for old data)
    for col in ['investor_score', 'balance_score', 'activity_score', 'defi_investor_score', 'reputation_score', 'sybil_risk_score', 'persona', 'wallet_type']:
        if col not in df.columns:
            if col == 'persona':
                df[col] = 'Unknown'
            elif col == 'wallet_type':
                df[col] = 'UNKNOWN'
            else:
                df[col] = 0.0

    # Sanctions columns (backward compat)
    if 'sanctions_hit' not in df.columns:
        df['sanctions_hit'] = False
    if 'sanctions_list_name' not in df.columns:
        df['sanctions_list_name'] = None
    if 'sanctions_entity_name' not in df.columns:
        df['sanctions_entity_name'] = None
    df['sanctions_hit'] = df['sanctions_hit'].fillna(False)

    # Intelligence columns (backward compat)
    if 'token_intelligence' not in df.columns:
        df['token_intelligence'] = None
    if 'persona_detail' not in df.columns:
        df['persona_detail'] = None
    if 'intent_signals' not in df.columns:
        df['intent_signals'] = None
    if 'staked_balances' not in df.columns:
        df['staked_balances'] = None
    if 'governance_balances' not in df.columns:
        df['governance_balances'] = None

    # Fill NaN scoring fields
    score_cols = ['investor_score', 'balance_score', 'activity_score', 'defi_investor_score', 'reputation_score', 'sybil_risk_score']
    for col in score_cols:
        df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0.0)
    df['persona'] = df['persona'].fillna('Unknown')
    df['wallet_type'] = df['wallet_type'].fillna('UNKNOWN')

    # Reorder columns for clarity
    cols_order = [
        'address', 'chain', 'tier', 'wallet_type', 'investor_score', 'persona',
        'est_net_worth_usd', 'stable_usd_total', 'native_balance',
        'balance_score', 'activity_score', 'defi_investor_score', 'reputation_score', 'sybil_risk_score',
        'product_relevance_score', 'tx_count', 'is_contract', 'known_entity_type', 'labels_str',
        'risk_flags_str', 'confidence', 'notes', 'last_scored_at', 'stable_balances',
        'activity_indicators', 'labels', 'risk_flags',
        'sanctions_hit', 'sanctions_list_name', 'sanctions_entity_name',
        'token_intelligence', 'persona_detail', 'intent_signals',
        'staked_balances', 'governance_balances',
    ]
    for col in cols_order:
        if col not in df.columns:
            df[col] = None

    return df[cols_order]


def generate_plot_base64(data, kind='pie', title='', colors=None):
    """Generates a base64-encoded PNG image of a plot."""
    plt.style.use('seaborn-v0_8-darkgrid')
    fig, ax = plt.subplots(figsize=(8, 5))

    if kind == 'pie':
        if len(data) == 0:
            plt.close(fig)
            return None
        wedge_colors = colors or plt.cm.Set3(np.linspace(0, 1, len(data)))
        data.plot(kind='pie', ax=ax, autopct='%1.1f%%', startangle=90, labels=None,
                  pctdistance=0.85, colors=wedge_colors)
        ax.legend(labels=data.index, loc='center left', bbox_to_anchor=(1, 0.5))
    elif kind == 'bar':
        if len(data) == 0:
            plt.close(fig)
            return None
        bar_colors = colors or ['#4A90D9'] * len(data)
        data.plot(kind='bar', ax=ax, color=bar_colors)
        ax.tick_params(axis='x', rotation=45)
    elif kind == 'hist':
        ax.hist(data, bins=20, color='#4A90D9', edgecolor='white', alpha=0.85)
        ax.set_xlabel('Score')
        ax.set_ylabel('Count')

    ax.set_title(title, fontsize=14, fontweight='bold')
    ax.set_ylabel('')
    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=120)
    plt.close(fig)
    return base64.b64encode(buf.getvalue()).decode('utf-8')


def _build_wealth_buckets(df: pd.DataFrame) -> list[dict]:
    """Group wallets into USD value buckets (users only)."""
    users = df[df['wallet_type'] == 'USER']
    buckets = [
        ("$0", 0, 1),
        ("$1 - $100", 1, 100),
        ("$100 - $1K", 100, 1_000),
        ("$1K - $10K", 1_000, 10_000),
        ("$10K - $100K", 10_000, 100_000),
        ("$100K - $1M", 100_000, 1_000_000),
        ("$1M+", 1_000_000, float('inf')),
    ]
    result = []
    for label, low, high in buckets:
        mask = (users['est_net_worth_usd'] >= low) & (users['est_net_worth_usd'] < high)
        subset = users[mask]
        if len(subset) > 0:
            result.append({
                "range": label,
                "count": len(subset),
                "total_usd": subset['est_net_worth_usd'].sum()
            })
    return result


def _find_labeled_wallets(df: pd.DataFrame) -> tuple[list[dict], dict]:
    """Find wallets matching known VC/KOL/smart_money labels."""
    labeled = []
    for _, row in df.iterrows():
        label = get_known_label(row['chain'], row['address'])
        if label and label['type'] in ('vc', 'kol', 'smart_money'):
            labeled.append({
                'address': row['address'],
                'chain': row['chain'],
                'label_name': label['name'],
                'label_type': label['type'],
                'est_net_worth_usd': row['est_net_worth_usd'] or 0,
                'investor_score': row['investor_score'] or 0,
            })

    vc_kol_usd = sum(w['est_net_worth_usd'] for w in labeled)
    summary = {
        'vc_kol_count': len(labeled),
        'vc_kol_usd': f"${vc_kol_usd:,.2f}",
    }
    return labeled, summary


def _wallet_type_breakdown(df: pd.DataFrame) -> list[dict]:
    """Build wallet_type summary table."""
    total = len(df)
    result = []
    for wt, group in df.groupby('wallet_type'):
        result.append({
            "wallet_type": wt,
            "count": len(group),
            "pct": f"{len(group) / total * 100:.1f}" if total > 0 else "0.0",
            "total_usd": group['est_net_worth_usd'].sum(),
        })
    # Sort: USER first, then alphabetical
    result.sort(key=lambda x: (0 if x['wallet_type'] == 'USER' else 1, x['wallet_type']))
    return result


def _aggregate_token_intelligence(df: pd.DataFrame) -> dict:
    """Aggregate token intelligence across all wallets."""
    total = len(df)
    if total == 0:
        return {"staking_count": 0, "governance_count": 0, "avg_stablecoin_share": 0, "avg_diversity": 0, "dry_powder_count": 0}

    staking_count = 0
    governance_count = 0
    stablecoin_shares = []
    diversities = []
    dry_powder_count = 0
    long_term_count = 0

    for _, row in df.iterrows():
        ti = row.get('token_intelligence')
        if not isinstance(ti, dict):
            continue
        if ti.get('has_staking_positions'):
            staking_count += 1
        if ti.get('has_governance_tokens'):
            governance_count += 1
        if ti.get('dry_powder_signal'):
            dry_powder_count += 1
        if ti.get('long_term_signal'):
            long_term_count += 1
        stablecoin_shares.append(ti.get('stablecoin_share', 0))
        diversities.append(ti.get('token_diversity', 0))

    return {
        "staking_count": staking_count,
        "staking_pct": f"{staking_count / total * 100:.1f}" if total > 0 else "0.0",
        "governance_count": governance_count,
        "governance_pct": f"{governance_count / total * 100:.1f}" if total > 0 else "0.0",
        "avg_stablecoin_share": f"{np.mean(stablecoin_shares):.1%}" if stablecoin_shares else "0.0%",
        "avg_diversity": f"{np.mean(diversities):.1f}" if diversities else "0",
        "dry_powder_count": dry_powder_count,
        "long_term_count": long_term_count,
    }


def _aggregate_intent_signals(df: pd.DataFrame) -> dict:
    """Aggregate investment intent signals across all wallets."""
    total = len(df)
    if total == 0:
        return {"readiness": {}, "signal_counts": {}, "total_deployable_usd": 0}

    readiness_counts = {"high": 0, "medium": 0, "low": 0}
    signal_counts = {}
    total_deployable = 0.0

    for _, row in df.iterrows():
        intent = row.get('intent_signals')
        if not isinstance(intent, dict):
            readiness_counts["low"] += 1
            continue
        r = intent.get('investment_readiness', 'low')
        readiness_counts[r] = readiness_counts.get(r, 0) + 1
        total_deployable += intent.get('estimated_deployable_usd', 0)
        for sig in intent.get('signals', []):
            name = sig.get('signal', 'unknown')
            signal_counts[name] = signal_counts.get(name, 0) + 1

    return {
        "readiness": readiness_counts,
        "signal_counts": dict(sorted(signal_counts.items(), key=lambda x: -x[1])),
        "total_deployable_usd": round(total_deployable, 2),
    }


def _aggregate_identity_data(df: pd.DataFrame) -> dict:
    """
    Aggregate identity enrichment data across all wallets.
    Extracts country/region signals, ENS adoption, NFT stats, and wallet age from
    activity_indicators dicts (where wallet_identity enrichment is stored).
    """
    total = len(df)
    if total == 0:
        return {"has_data": False}

    country_counts: dict[str, int] = {}
    region_counts: dict[str, int] = {}
    funding_source_counts: dict[str, int] = {}
    ens_count = 0
    nft_holder_count = 0
    blue_chip_count = 0
    blue_chip_collections: dict[str, int] = {}
    wallet_ages: list[int] = []
    ens_wallets: list[dict] = []

    for _, row in df.iterrows():
        ai = row.get("activity_indicators")
        if not isinstance(ai, dict):
            continue

        # Country / region
        country = ai.get("country_hint")
        if country and country != "Global":
            country_counts[country] = country_counts.get(country, 0) + 1
        region = ai.get("region")
        if region and region != "Global":
            region_counts[region] = region_counts.get(region, 0) + 1

        # Funding source CEX
        fs = ai.get("funding_source")
        if fs:
            funding_source_counts[fs] = funding_source_counts.get(fs, 0) + 1

        # ENS names
        ens_name = ai.get("ens_name")
        if ens_name:
            ens_count += 1
            ens_wallets.append({
                "address": row.get("address", ""),
                "chain": row.get("chain", ""),
                "ens_name": ens_name,
                "est_net_worth_usd": row.get("est_net_worth_usd", 0),
                "tier": row.get("tier", ""),
            })

        # NFTs
        if ai.get("has_nfts"):
            nft_holder_count += 1
        if ai.get("has_blue_chip_nfts"):
            blue_chip_count += 1
            for col in ai.get("blue_chip_collections", []):
                blue_chip_collections[col] = blue_chip_collections.get(col, 0) + 1

        # Wallet age
        age = ai.get("wallet_age_days")
        if isinstance(age, int) and age >= 0:
            wallet_ages.append(age)

    # Sort by descending count
    country_breakdown = sorted(country_counts.items(), key=lambda x: -x[1])
    region_breakdown = sorted(region_counts.items(), key=lambda x: -x[1])
    funding_breakdown = sorted(funding_source_counts.items(), key=lambda x: -x[1])
    blue_chip_breakdown = sorted(blue_chip_collections.items(), key=lambda x: -x[1])
    ens_wallets_top = sorted(ens_wallets, key=lambda x: -x["est_net_worth_usd"])[:20]

    # Wallet age buckets
    age_buckets = _build_wallet_age_buckets(wallet_ages)

    enriched_count = sum(1 for _, row in df.iterrows()
                         if isinstance(row.get("activity_indicators"), dict) and
                         row["activity_indicators"].get("wallet_age_days") is not None)

    return {
        "has_data": enriched_count > 0,
        "enriched_count": enriched_count,
        "enriched_pct": f"{enriched_count / total * 100:.1f}" if total > 0 else "0.0",
        # Country
        "country_breakdown": country_breakdown,
        "region_breakdown": region_breakdown,
        "country_identified_count": sum(v for _, v in country_counts.items()),
        # Funding sources
        "funding_breakdown": funding_breakdown,
        # ENS
        "ens_count": ens_count,
        "ens_pct": f"{ens_count / total * 100:.1f}" if total > 0 else "0.0",
        "ens_wallets": ens_wallets_top,
        # NFTs
        "nft_holder_count": nft_holder_count,
        "nft_holder_pct": f"{nft_holder_count / total * 100:.1f}" if total > 0 else "0.0",
        "blue_chip_count": blue_chip_count,
        "blue_chip_breakdown": blue_chip_breakdown,
        # Wallet age
        "avg_wallet_age_days": int(sum(wallet_ages) / len(wallet_ages)) if wallet_ages else None,
        "wallet_age_count": len(wallet_ages),
        "age_buckets": age_buckets,
    }


def _build_wallet_age_buckets(ages: list[int]) -> list[dict]:
    """Group wallet ages into human-readable buckets."""
    buckets = [
        ("< 6 months",  0,    180),
        ("6–12 months", 180,  365),
        ("1–2 years",   365,  730),
        ("2–3 years",   730,  1095),
        ("3+ years",    1095, 99999),
    ]
    result = []
    for label, lo, hi in buckets:
        count = sum(1 for a in ages if lo <= a < hi)
        if count > 0:
            result.append({"range": label, "count": count})
    return result


def _build_persona_detail_table(df: pd.DataFrame) -> list[dict]:
    """Build persona detail table with confidence from persona_detail JSON."""
    rows = []
    for _, row in df.iterrows():
        pd_data = row.get('persona_detail')
        if isinstance(pd_data, dict):
            rows.append({
                'address': row['address'],
                'chain': row['chain'],
                'persona': pd_data.get('primary', row.get('persona', 'Unknown')),
                'confidence': pd_data.get('confidence', 0.5),
                'evidence': '; '.join(pd_data.get('evidence', [])),
                'secondary': pd_data.get('secondary') or '-',
                'est_net_worth_usd': row.get('est_net_worth_usd', 0),
            })
    # Sort by confidence descending, take top 20
    rows.sort(key=lambda x: -x['confidence'])
    return rows[:20]


def _aggregate_cross_chain(df: pd.DataFrame) -> pd.DataFrame:
    """
    When the same address was scanned on multiple chains, aggregate into one row per address:
    - Sum financial values (USD net worth, stablecoins, tx count)
    - Re-compute balance_score and investor_score from combined USD
    - Upgrade tier if combined value crosses a higher threshold
    - Add 'chains_with_assets' column listing chains where the wallet has holdings
    - Add 'scan_chains' column listing all chains that were scanned for this address

    Single-chain wallets pass through unchanged (only gains the two new columns).
    """
    if df.empty:
        return df

    addr_counts = df['address'].value_counts()
    multi_addrs = set(addr_counts[addr_counts > 1].index)

    # Add columns for all rows; single-chain wallets are simple
    df = df.copy()
    df['scan_chains'] = df['chain']
    df['chains_with_assets'] = df.apply(
        lambda r: r['chain'] if float(r.get('est_net_worth_usd', 0) or 0) > 0 else '', axis=1
    )

    if not multi_addrs:
        return df

    single_df = df[~df['address'].isin(multi_addrs)].copy()
    multi_df = df[df['address'].isin(multi_addrs)]

    aggregated = []
    for address, group in multi_df.groupby('address'):
        # Base row: chain with highest USD value
        base = group.loc[group['est_net_worth_usd'].idxmax()].copy()

        # Chains that were scanned
        base['scan_chains'] = ', '.join(sorted(group['chain'].tolist()))

        # Chains where there are actual holdings
        has_value = group[group['est_net_worth_usd'].fillna(0) > 0]['chain'].tolist()
        base['chains_with_assets'] = ', '.join(sorted(set(has_value))) if has_value else ''

        # Sum financials
        combined_usd = float(group['est_net_worth_usd'].fillna(0).sum())
        combined_stable = float(group['stable_usd_total'].fillna(0).sum())
        combined_tx = int(group['tx_count'].fillna(0).sum())
        base['est_net_worth_usd'] = combined_usd
        base['stable_usd_total'] = combined_stable
        base['tx_count'] = combined_tx

        # Re-compute balance_score from combined USD (same formula as scoring.py)
        total_for_score = combined_usd + combined_stable
        new_balance_score = min(round(math.log10(max(total_for_score, 1)) * 10, 1), 100.0) if total_for_score > 0 else 0.0
        base['balance_score'] = new_balance_score

        # Re-compute investor_score (keep other component scores from base chain)
        w = {"balance": 0.30, "activity": 0.15, "defi": 0.25, "reputation": 0.20, "sybil": -0.10}
        new_investor_score = min(max(round(
            new_balance_score * w["balance"]
            + float(base.get('activity_score', 0) or 0) * w["activity"]
            + float(base.get('defi_investor_score', 0) or 0) * w["defi"]
            + float(base.get('reputation_score', 0) or 0) * w["reputation"]
            + float(base.get('sybil_risk_score', 0) or 0) * w["sybil"]
        , 1), 0.0), 100.0)
        base['investor_score'] = new_investor_score

        # Re-determine tier if this is a USER wallet
        if base.get('wallet_type') == 'USER':
            if new_investor_score >= 55:
                base['tier'] = 'Whale'
            elif new_investor_score >= 30:
                base['tier'] = 'Tuna'
            else:
                base['tier'] = 'Fish'

        # If ANY chain classified this as CEX/infra, use that type
        wt_priority = ['CEX_EXCHANGE', 'DEX_ROUTER', 'BRIDGE', 'PROTOCOL', 'CONTRACT', 'UNKNOWN', 'USER']
        for wt in wt_priority:
            if wt in group['wallet_type'].values:
                base['wallet_type'] = wt
                break

        aggregated.append(base)

    agg_df = pd.DataFrame(aggregated)
    return pd.concat([single_df, agg_df], ignore_index=True)


def generate_executive_report(df: pd.DataFrame, job_id: int, project_name: str = "Project", output_format='markdown', total_wallets_actual: int | None = None, chains_scanned: int | None = None) -> str:
    """Generates the full executive intelligence report with narrative-first structure."""

    # Aggregate cross-chain scans: one row per unique address with combined financials
    scans_count = len(df)  # total individual chain scans in this sample
    df = _aggregate_cross_chain(df)
    total_wallets = total_wallets_actual or len(df)  # unique addresses (not chain-scan count)
    chains_count = chains_scanned or (1 if scans_count == len(df) else scans_count // max(len(df), 1))
    is_multi_chain = chains_scanned is not None and chains_scanned > 1
    report_is_sample = total_wallets_actual is not None and total_wallets_actual > len(df)
    failed_count = int(df['notes'].fillna('').str.startswith('Analysis failed:').sum()) if 'notes' in df.columns else 0
    tier_dist = df['tier'].value_counts()
    total_usd = df['est_net_worth_usd'].sum()
    chain_dist = df['chain'].value_counts()
    persona_dist = df['persona'].value_counts()

    # Wallet type counts
    wt_counts = df['wallet_type'].value_counts()
    user_count = int(wt_counts.get('USER', 0))
    cex_count = int(wt_counts.get('CEX_EXCHANGE', 0))
    dex_count = int(wt_counts.get('DEX_ROUTER', 0))
    bridge_count = int(wt_counts.get('BRIDGE', 0))
    protocol_count = int(wt_counts.get('PROTOCOL', 0))
    contract_count = int(wt_counts.get('CONTRACT', 0))
    infra_total = total_wallets - user_count
    user_pct = f"{user_count / total_wallets * 100:.1f}" if total_wallets > 0 else "0.0"

    # Tier counts (among users only)
    users_df = df[df['wallet_type'] == 'USER']
    whale_count = int((users_df['tier'] == 'Whale').sum())
    tuna_count = int((users_df['tier'] == 'Tuna').sum())
    fish_count = int((users_df['tier'] == 'Fish').sum())
    whale_usd_val = users_df[users_df['tier'] == 'Whale']['est_net_worth_usd'].sum()
    whale_usd = f"${whale_usd_val:,.2f}"

    # Scoring averages (users only, or all if no users)
    score_df = users_df if len(users_df) > 0 else df
    score_cols = ['balance_score', 'activity_score', 'defi_investor_score', 'reputation_score', 'sybil_risk_score']
    score_averages = {col: f"{score_df[col].mean():.1f}" for col in score_cols}
    avg_investor_score = f"{score_df['investor_score'].mean():.1f}"

    # Risk
    concentration_whales = df[df['tier'] == 'Whale']
    concentration_risk = concentration_whales['est_net_worth_usd'].sum() / total_usd if total_usd > 0 else 0
    sybil_risk_wallets = df[df['sybil_risk_score'] >= 40]

    # Wealth buckets (users only)
    wealth_buckets = _build_wealth_buckets(df)

    # Labeled wallets
    labeled_wallets, labeled_summary = _find_labeled_wallets(df)

    # Wallet type breakdown table
    wallet_type_breakdown = _wallet_type_breakdown(df)

    # Sanctions aggregation
    sanctioned_df = df[df['sanctions_hit'] == True]
    sanctions_count = len(sanctioned_df)
    sanctions_wallets = []
    if sanctions_count > 0:
        sanctions_wallets = sanctioned_df[['address', 'chain', 'sanctions_list_name', 'sanctions_entity_name', 'est_net_worth_usd']].rename(
            columns={'sanctions_list_name': 'list_name', 'sanctions_entity_name': 'entity_name'}
        ).to_dict('records')
    sanctions_enabled = True  # always show section if data exists

    # --- Project Health & Community Quality ---
    community_score = project_health.compute_community_quality_score(df)
    health_flags = project_health.compute_health_flags(df)
    concentration_metrics = project_health.compute_concentration_metrics(df)

    # --- Token Intelligence Aggregation ---
    token_intel_agg = _aggregate_token_intelligence(df)

    # --- Intent Signals Aggregation ---
    intent_agg = _aggregate_intent_signals(df)

    # --- Identity & Geographic Intelligence ---
    identity_agg = _aggregate_identity_data(df)

    # --- Persona Detail Table ---
    persona_detail_table = _build_persona_detail_table(df)

    # Tables — top by score (users only)
    top_by_score_df = users_df.sort_values('investor_score', ascending=False).head(20)
    top_whales = df.sort_values('est_net_worth_usd', ascending=False).head(20)
    high_activity_low_balance = df[(df['tx_count'] > 50) & (df['est_net_worth_usd'] < 1000)].head(20)

    # Infra-type tables
    cex_wallets = df[df['wallet_type'] == 'CEX_EXCHANGE'].sort_values('est_net_worth_usd', ascending=False)
    dex_wallets = df[df['wallet_type'] == 'DEX_ROUTER']
    bridge_wallets = df[df['wallet_type'] == 'BRIDGE'].sort_values('est_net_worth_usd', ascending=False)

    # Chain breakdown with avg score
    chain_breakdown_df = df.groupby('chain').agg(
        count=('est_net_worth_usd', 'count'),
        sum=('est_net_worth_usd', 'sum'),
        avg_score=('investor_score', 'mean')
    ).reset_index()

    # Stablecoin vs native breakdown
    total_stable = df['stable_usd_total'].sum()
    total_native = total_usd - total_stable
    stable_vs_native = pd.Series({'Stablecoin': max(total_stable, 0), 'Native Token': max(total_native, 0)})

    # Generate plots
    tier_colors = {'Whale': '#1e40af', 'Tuna': '#0891b2', 'Fish': '#65a30d', 'Infra': '#9ca3af'}
    tier_color_list = [tier_colors.get(t, '#6b7280') for t in tier_dist.index]

    plots = {
        "tier_dist_pie": generate_plot_base64(tier_dist, kind='pie', title='Wallet Tier Distribution', colors=tier_color_list),
        "chain_dist_bar": generate_plot_base64(chain_dist, kind='bar', title='Wallets by Chain'),
        "investor_score_hist": generate_plot_base64(score_df['investor_score'].values, kind='hist', title='Investor Score Distribution'),
        "stable_vs_native_pie": generate_plot_base64(stable_vs_native, kind='pie', title='Stablecoin vs Native Holdings'),
    }

    # Persona distribution pie chart
    if len(persona_dist) > 0:
        plots["persona_dist_pie"] = generate_plot_base64(persona_dist, kind='pie', title='Persona Distribution')
    else:
        plots["persona_dist_pie"] = None

    # Extract ENS name from activity_indicators into a display column
    def _get_ens(row):
        ai = row.get('activity_indicators')
        if isinstance(ai, dict):
            return ai.get('ens_name') or ''
        return ''
    df['ens_name'] = df.apply(_get_ens, axis=1)

    def _get_country(row):
        ai = row.get('activity_indicators')
        if isinstance(ai, dict):
            return ai.get('country_hint') or ''
        return ''
    df['country_hint'] = df.apply(_get_country, axis=1)

    def _get_wallet_age(row):
        ai = row.get('activity_indicators')
        if isinstance(ai, dict):
            age = ai.get('wallet_age_days')
            if age is not None:
                return f"{age // 365}y {(age % 365) // 30}m" if age >= 365 else f"{age // 30}m"
        return ''
    df['wallet_age'] = df.apply(_get_wallet_age, axis=1)

    _available = set(df.columns.tolist())
    def _cols(*cols): return [c for c in cols if c in _available]
    table_cols_score = _cols('address', 'chains_with_assets', 'tier', 'investor_score', 'persona', 'est_net_worth_usd', 'wallet_type', 'ens_name')
    table_cols_whales = _cols('address', 'chains_with_assets', 'tier', 'est_net_worth_usd', 'stable_usd_total', 'wallet_type', 'labels_str', 'ens_name')
    table_cols_infra = _cols('address', 'chain', 'labels_str', 'est_net_worth_usd', 'confidence')
    table_cols_sybil = _cols('address', 'chain', 'tx_count', 'est_net_worth_usd', 'sybil_risk_score', 'wallet_type')

    context = {
        "project_name": project_name,
        "job_id": job_id,
        "reference_id": generate_reference_id(job_id),
        "generation_date": pd.Timestamp.now(tz='UTC').strftime('%Y-%m-%d %H:%M:%S UTC'),
        "total_wallets": total_wallets,
        "chains_scanned": chains_count,
        "scans_count": scans_count,
        "is_multi_chain": is_multi_chain,
        "failed_count": failed_count,
        "total_usd_controlled": f"${total_usd:,.2f}",
        "chain_count": len(chain_dist),
        # Wallet type counts
        "user_count": user_count,
        "user_pct": user_pct,
        "infra_total": infra_total,
        "cex_count": cex_count,
        "dex_count": dex_count,
        "bridge_count": bridge_count,
        "protocol_count": protocol_count,
        "contract_count": contract_count,
        # Tier counts (users only)
        "whale_count": whale_count,
        "tuna_count": tuna_count,
        "fish_count": fish_count,
        "whale_usd": whale_usd,
        # Scoring
        "avg_investor_score": avg_investor_score,
        "score_averages": score_averages,
        "persona_distribution": persona_dist.to_dict(),
        "wealth_buckets": wealth_buckets,
        "labeled_wallets": labeled_wallets,
        "labeled_summary": labeled_summary,
        "wallet_type_breakdown": wallet_type_breakdown,
        "risk_overview": {
            "concentration_pct": f"{concentration_risk:.2%}",
            "num_whales": whale_count,
            "sybil_risk_count": len(sybil_risk_wallets),
        },
        "tables": {
            "top_by_score": top_by_score_df[[c for c in table_cols_score if c in top_by_score_df.columns]].to_dict('records') if len(top_by_score_df) > 0 else [],
            "top_whales": top_whales[[c for c in table_cols_whales if c in top_whales.columns]].to_dict('records') if len(top_whales) > 0 else [],
            "high_activity_low_balance": high_activity_low_balance[table_cols_sybil].to_dict('records') if len(high_activity_low_balance) > 0 else [],
            "chain_breakdown": chain_breakdown_df.to_dict('records'),
            "failed_wallets": df[df['notes'].fillna('').str.startswith('Analysis failed:')][['address', 'chain', 'notes']].to_dict('records') if failed_count > 0 else [],
            "cex_wallets": cex_wallets[table_cols_infra].to_dict('records') if len(cex_wallets) > 0 else [],
            "dex_wallets": dex_wallets[['address', 'chain', 'labels_str', 'confidence']].to_dict('records') if len(dex_wallets) > 0 else [],
            "bridge_wallets": bridge_wallets[table_cols_infra].to_dict('records') if len(bridge_wallets) > 0 else [],
            "persona_detail": persona_detail_table,
        },
        "plots": plots,
        # Sanctions
        "sanctions_enabled": sanctions_enabled,
        "sanctions_count": sanctions_count,
        "sanctions_wallets": sanctions_wallets,
        # Community Quality Score
        "community_score": community_score,
        # Health Flags
        "health_flags": health_flags,
        # Concentration Metrics
        "concentration_metrics": concentration_metrics,
        # Token Intelligence Aggregation
        "token_intel": token_intel_agg,
        # Intent Signals Aggregation
        "intent_agg": intent_agg,
        # Identity & Geographic Intelligence
        "identity_agg": identity_agg,
    }

    template = jinja_env.get_template('report_template.md')
    md_report = template.render(context)

    if output_format == 'html':
        html_report = markdown2.markdown(md_report, extras=["tables", "fenced-code-blocks"])
        return _wrap_html(html_report, project_name)

    return md_report


def _wrap_html(html_body: str, project_name: str) -> str:
    """Wraps HTML report body with full document structure and CSS for both web and PDF."""
    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{project_name} — Wallet Intelligence Report</title>
    <style>
        @page {{ size: A4; margin: 15mm; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; padding: 20px; max-width: 1000px; margin: auto; color: #1e293b; font-size: 14px; }}
        h1 {{ color: #0f172a; font-size: 24px; border-bottom: 3px solid #3b82f6; padding-bottom: 8px; }}
        h2 {{ color: #1e40af; font-size: 20px; margin-top: 30px; border-bottom: 1px solid #e2e8f0; padding-bottom: 4px; }}
        h3 {{ color: #334155; font-size: 16px; margin-top: 20px; }}
        table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; font-size: 13px; }}
        th {{ background-color: #1e40af; color: white; padding: 8px 10px; text-align: left; font-weight: 600; }}
        td {{ border: 1px solid #e2e8f0; padding: 6px 10px; }}
        tr:nth-child(even) {{ background-color: #f8fafc; }}
        tr:hover {{ background-color: #eff6ff; }}
        img {{ max-width: 100%; height: auto; display: block; margin: 15px auto; }}
        hr {{ border: none; border-top: 2px solid #e2e8f0; margin: 25px 0; }}
        code {{ background-color: #f1f5f9; padding: 1px 5px; border-radius: 3px; font-size: 12px; word-break: break-all; }}
        strong {{ color: #0f172a; }}
        p {{ margin-bottom: 10px; }}
        ul {{ margin-bottom: 15px; }}
        li {{ margin-bottom: 5px; }}
    </style>
</head>
<body>{html_body}</body>
</html>"""


def generate_pdf(html_content: str) -> bytes:
    """Converts HTML report to PDF using WeasyPrint."""
    from weasyprint import HTML
    return HTML(string=html_content).write_pdf()


def generate_docx(html_content: str, project_name: str) -> bytes:
    """Converts HTML report to Word document using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import re
    from bs4 import BeautifulSoup
    import io

    doc = Document()

    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    body = soup.find('body') or soup

    # Style settings
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for element in body.children:
        if not hasattr(element, 'name') or element.name is None:
            continue

        text = element.get_text(strip=True) if hasattr(element, 'get_text') else str(element)

        if element.name == 'h1':
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(15, 23, 42)
        elif element.name == 'h2':
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        elif element.name == 'h3':
            doc.add_heading(text, level=3)
        elif element.name == 'p':
            doc.add_paragraph(text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(strip=True), style='List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                # Count columns from header
                header_cells = rows[0].find_all(['th', 'td'])
                num_cols = len(header_cells)
                if num_cols > 0:
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for row_idx, tr in enumerate(rows):
                        cells = tr.find_all(['th', 'td'])
                        for col_idx, cell in enumerate(cells):
                            if col_idx < num_cols:
                                table_cell = table.rows[row_idx].cells[col_idx]
                                table_cell.text = cell.get_text(strip=True)
                                # Bold headers
                                if cell.name == 'th':
                                    for paragraph in table_cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

                    doc.add_paragraph()  # spacing after table
        elif element.name == 'hr':
            doc.add_paragraph('─' * 50)
        elif element.name == 'img':
            # Skip images in Word for now (base64 charts are complex)
            doc.add_paragraph('[Chart - see HTML/PDF version]')

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
